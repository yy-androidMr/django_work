import docx
from win32com import client as wc

import os

ch_equal = '＝'
ch_add = '＋'
ch_reduce = '－'


def convert_all2en(str):
    str = str.replace(ch_equal, '=')
    str = str.replace(ch_reduce, '-')
    str = str.replace(ch_add, '+')
    return str


def convert_all2ch(str):
    str = str.replace('+', ch_add)
    str = str.replace('=', ch_equal)
    str = str.replace('-', ch_reduce)
    return str


def is_Chinese(word):
    for ch in word:
        if '\u4e00' <= ch <= '\u9fff':
            return True
    return False


def calculation(eval_str):
    anser = eval(eval_str)
    # print(eval_str + '=' + str(anser))
    return anser


def run(run, content, anser_pool):
    # print(run.text)
    content = convert_all2en(content)
    # split_content = re.split('＝|=', content)
    split_content = content.split('=')
    re_content = ''
    for item in split_content:
        ca = item.strip()
        re_content += item
        if ca is None or ca == '':
            pass
        else:
            value = calculation(ca)
            anser_pool.append(value)
            re_content += '=' + str(value)
            # print(re_content)
            # 这里尝试重组
            # re_content ='＝'.join(split_content)
    print(content + "\n" + re_content)
    run.text = convert_all2ch(re_content)


def save_result(runs_map):
    for item_list in runs_map:
        print('-------')
        for item in item_list:
            print(item.text)
    pass


def org_ansers(list):
    write_list = []
    for item in list:
        str = '; '.join('%s' % id for id in item)
        write_list.append(str)
        print(str)
        # for item in list:
        #     print(item)
        # pass
    return write_list


def read_oral(path):
    file = docx.Document(path)
    start_str = '口算'

    start_calculation = False
    anser_pool = []
    anser_list = []
    for i in range(len(file.paragraphs)):
        # print(file.paragraphs[i])
        content = file.paragraphs[i].text

        if start_calculation and is_Chinese(content):
            start_calculation = False
            anser_pool.append(anser_list)
            print('停止记录')

        if start_calculation:
            if len(file.paragraphs[i].runs) > 0:
                runs_content = ''
                for item in file.paragraphs[i].runs:
                    runs_content += item.text
                run(file.paragraphs[i].runs[0], runs_content, anser_list)

        if start_str in content:
            start_calculation = True
            anser_list = []
            print('开始记录')

            # for line in file.paragraphs:
            #     print(line.text)

    # save_result(runs_list)
    return org_ansers(anser_pool)
    # file.save(r'G:\cache\audio\3.docx')


# read_oral(r"G:\cache\audio\【计算专项课程】A3 第3讲（作业版）.docx")

# a='\u2002'
# print('sadfasdf:'+a)


def doSaveAas(file_name):
    new_file = file_name + '.docx'
    if not os.path.exists(new_file):
        word = wc.Dispatch('Word.Application')
        doc = word.Documents.Open(file_name)  # 目标路径下的文件
        doc.SaveAs(new_file, 12, False, "", True, "", False, False, False,
                   False)  # 转化后路径下的文件
        doc.Close()
        word.Quit()
    return new_file


dir_path = r'G:\cache\audio'


def begin():
    docx_dict = {}
    for root, dirs, files in os.walk(dir_path):
        for file in files:
            if '.doc' in file and '.docx' not in file:
                source_path = os.path.join(root, file).replace('\\', '/')
                docx_dict[source_path] = doSaveAas(source_path)
                # docx_dict.append(source_path)
                # print(source_path)

    f = open(dir_path + '/out.txt', "w+")
    f.write('口算题')
    for k, v in docx_dict.items():
        print(k + "   :" + v)
        result = read_oral(v)
        f.write('\n')
        # for docx in docx_list:
        #     print(docx)
        f.write(os.path.basename(k) + '\n')
        for i in result:
            f.write(str(i) + '\n')  # \r\n为换行符  # doSaveAas('【计算专项课程】A3 第2讲（作业版）')
        os.remove(v)
    f.close()


# doSaveAas('【计算专项课程】A3 第3讲（作业版）')

begin()
